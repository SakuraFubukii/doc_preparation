"""
Word文档转换器模块

此模块提供Word文档(.docx)转换为Markdown文件的功能。
"""

import json
from docx import Document
from pathlib import Path
from tqdm import tqdm

import config
from core.utils import (
    clean_markdown, extract_metadata, save_images,
    ensure_dir, is_temp_file, is_short_text,
    combine_text_fragments, write_json_file,
    post_process_markdown_content
)

def extract_table_data(table):
    """提取表格数据"""
    headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
    data_rows = []
    
    for row in table.rows[1:]:
        row_data = {}
        for i, cell in enumerate(row.cells):
            key = headers[i] if i < len(headers) else f"列{i+1}"
            row_data[key] = cell.text.strip()
        if row_data:
            data_rows.append(row_data)
    
    return headers, data_rows

def process_paragraph(para, pending_short_text):
    """处理段落文本"""
    text = para.text.strip()
    
    # 处理标题
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 2
        return f"{'#' * level} {text}\n\n", []
    
    # 处理列表
    if 'List' in para.style.name:
        prefix = '1. ' if 'Number' in para.style.name else '- '
        return f"{prefix}{text}\n", []
    
    # 处理短文本
    if is_short_text(text, config.SHORT_TEXT_THRESHOLD):
        return "", pending_short_text + [text]
    
    # 处理长文本
    return combine_text_fragments(pending_short_text + [text]) + "\n\n", []



def save_table_as_json_txt(table_data, output_folder, table_index, file_base_name):
    """保存表格数据为JSON格式的txt文件"""
    output_folder = Path(output_folder)
    
    # 创建表格文件名
    table_filename = f"{file_base_name}_table_{table_index + 1}.txt"
    table_file_path = output_folder / table_filename
    
    # 保存表格数据为JSON格式
    try:
        with open(table_file_path, 'w', encoding='utf-8') as f:
            json.dump(table_data, f, ensure_ascii=False, indent=2)
        print(f"  - 已生成表格文件: {table_filename}")
        return table_file_path
    except Exception as e:
        print(f"  - 保存表格文件失败: {str(e)}")
        return None

def convert_docx_to_markdown(docx_path, output_folder):
    """将Word文档转换为Markdown格式"""
    docx_path, output_folder = Path(docx_path), Path(output_folder)
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"无法打开文档 {docx_path}: {str(e)}")
        return "", [], {}
    
    title = getattr(doc.core_properties, 'title', docx_path.stem)
    md_lines = [f"# {title}\n\n"]
    tables_data = []
    pending_short_text = []
    metadata = extract_metadata(doc)
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            para = next((p for p in doc.paragraphs if p._element is element), None)
            if para:
                content, pending_short_text = process_paragraph(para, pending_short_text)
                if content:
                    md_lines.append(content)
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            table = doc.tables[len(tables_data)]
            # 处理待处理的短文本
            if pending_short_text:
                md_lines.append(combine_text_fragments(pending_short_text) + "\n\n")
                pending_short_text = []
            
            # 表格处理：使用纯文本
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    md_lines.append(f"| {row_text} |\n")
            md_lines.append("\n")
            
            # 提取表格数据
            headers, data_rows = extract_table_data(table)
            table_data = {"headers": headers, "data": data_rows}
            tables_data.append(table_data)
            
            # 为每个表格生成单独的JSON txt文件
            save_table_as_json_txt(table_data, output_folder, len(tables_data) - 1, docx_path.stem)
    
    # 处理剩余的短文本
    if pending_short_text:
        md_lines.append(combine_text_fragments(pending_short_text) + "\n\n")
    
    # 添加图片
    for img_path in save_images(doc, output_folder):
        md_lines.append(f"![图片]({img_path})\n\n")
    
    # 应用完整的文本清洗和标准化处理
    raw_markdown = "".join(md_lines)
    cleaned_markdown = post_process_markdown_content(raw_markdown)
    
    return cleaned_markdown, tables_data, metadata

def process_single_docx(input_path, output_folder):
    """处理单个Word文档"""
    input_path, output_folder = Path(input_path), Path(output_folder)
    
    if is_temp_file(input_path):
        print(f"⚠️ 跳过临时文件: {input_path.name}")
        return False
    
    base_name = input_path.stem
    output_subfolder = ensure_dir(output_folder / base_name)
    
    try:
        md_content, tables_data, metadata = convert_docx_to_markdown(input_path, output_subfolder)
        
        # 保存文件
        (output_subfolder / f"{base_name}.md").write_text(md_content, encoding='utf-8')
        write_json_file(metadata, output_subfolder / f"{base_name}_metadata.json")
        
        if tables_data:
            write_json_file(tables_data, output_subfolder / f"{base_name}_tables.json")
        
        print(f"✅ 成功转换: {input_path.name}")
        print(f"  输出目录: {output_subfolder}")
        return True
    except Exception as e:
        print(f"❌ 处理失败: {input_path.name} - {str(e)}")
        return False

def convert_batch(input_folder, output_folder, file_extensions=None):
    """批量转换Word文档"""
    file_extensions = file_extensions or ['.docx', '.doc']
    input_path, output_folder = Path(input_folder), Path(output_folder)
    
    if not input_path.exists():
        print(f"错误: 输入文件夹 '{input_path}' 不存在")
        return 0, 0
    
    ensure_dir(output_folder)
    doc_files = [f for f in input_path.iterdir() 
                if str(f.suffix).lower() in file_extensions and not is_temp_file(f)]
    
    if not doc_files:
        print(f"在 {input_folder} 中没有找到Word文档")
        return 0, 0
    
    print(f"找到 {len(doc_files)} 个Word文档，开始转换...")
    success_count = 0
    
    for doc_file in tqdm(doc_files, desc="转换文档"):
        if process_single_docx(doc_file, output_folder):
            success_count += 1

    print("\n" + "=" * 50)
    print(f"处理完成! 成功转换: {success_count}/{len(doc_files)}")
    print(f"输出目录: {Path(output_folder).absolute()}")
    print("=" * 50)
    
    return success_count, len(doc_files)


def main():
    """作为独立模块运行时的入口点"""
    import sys
    import os
    
    # 添加项目根目录到路径
    sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
    
    # 默认配置
    INPUT_FOLDER = "input"
    OUTPUT_FOLDER = "output"
    
    # 解析命令行参数
    if len(sys.argv) > 1:
        INPUT_FOLDER = sys.argv[1]
    if len(sys.argv) > 2:
        OUTPUT_FOLDER = sys.argv[2]
    
    convert_batch(INPUT_FOLDER, OUTPUT_FOLDER)


if __name__ == "__main__":
    main()